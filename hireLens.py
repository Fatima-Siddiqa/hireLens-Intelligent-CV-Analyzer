import streamlit as st
from PIL import Image
import base64
import zipfile
import io
import time
from PyPDF2 import PdfReader
import docx
import pandas as pd
import math
import re

def prep_kw(kw):
    s = kw.strip()
    if s.endswith("s"): return s, s[:-1]  # plural + singular
    else: return s+"s", s  # singular + plural

def build_patterns(keywords):
    patterns = []
    for kw in keywords:
        kw_lower = kw.lower()
        p1, p2 = prep_kw(kw_lower)
        # whole word only
        patterns.append(re.compile(rf"\b{p1}\b", re.IGNORECASE))
        patterns.append(re.compile(rf"\b{p2}\b", re.IGNORECASE))
    return patterns

def count_matches(text, patterns):
    text_lower = text.lower()
    total = 0
    for p in patterns:
        total += len(p.findall(text_lower))
    return total

# Helpers
def img_to_base64(path):
    with open(path, "rb") as f:
        return base64.b64encode(f.read()).decode()

# Text extraction functions 
def extract_text_from_pdf(file_like):
    try:
        reader = PdfReader(file_like)
    except Exception:
        # If PdfReader fails on zipfile's ZipExtFile directly, wrap bytes
        file_like = io.BytesIO(file_like.read())
        reader = PdfReader(file_like)
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text

def extract_text_from_docx(file_like):
    # docx.Document accepts a file-like object
    try:
        doc = docx.Document(file_like)
    except Exception:
        # if that fails, read bytes into BytesIO
        file_like = io.BytesIO(file_like.read())
        doc = docx.Document(file_like)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

# Pattern search implementations 
# Each returns: (occurrences_count, comparisons_count, elapsed_ms)

# BRUTE FORCE
def search_naive(text, pattern):
    if not pattern:
        return 0, 0, 0.0

    n = len(text)
    m = len(pattern)
    comparisons = 0
    occ = 0

    start = time.perf_counter()

    for i in range(0, n-m+1):
        ok = True
        for j in range(m):
            comparisons += 1
            if text[i+j] != pattern[j]:
                ok = False
                break
        if ok:
            occ += 1

    elapsed = (time.perf_counter() - start)*1000.0
    return occ, comparisons, elapsed


# RABIN-KARP
def search_rabin_karp(text, pattern):
    n = len(text)
    m = len(pattern)
    if m == 0 or n < m:
        return 0, 0, 0.0

    base = 256
    mod  = 101
    comparisons = 0

    start = time.perf_counter()

    h = pow(base, m-1, mod)

    pat_hash = 0
    txt_hash = 0
    for i in range(m):
        pat_hash = (pat_hash*base + ord(pattern[i])) % mod
        txt_hash = (txt_hash*base + ord(text[i]))    % mod

    occ = 0

    for i in range(n-m+1):
        # hash match then actual compare
        if pat_hash == txt_hash:
            comparisons += 1
            if text[i:i+m] == pattern:
                occ += 1

        # slide window
        if i < n-m:
            txt_hash = (txt_hash - ord(text[i])*h)*base + ord(text[i+m])
            txt_hash %= mod

    elapsed = (time.perf_counter() - start)*1000.0
    return occ, comparisons, elapsed

def build_kmp_lps(p):
    m = len(p)
    lps = [0]*m
    length = 0
    i = 1
    while i < m:
        if p[i] == p[length]:
            length += 1
            lps[i] = length
            i += 1
        else:
            if length != 0:
                length = lps[length-1]
            else:
                lps[i] = 0
                i += 1
    return lps

def search_kmp(text, pattern):
    if not pattern:
        return 0, 0, 0.0
    n, m = len(text), len(pattern)
    lps = build_kmp_lps(pattern)
    i = j = 0
    comparisons = 0
    occ = 0
    start = time.perf_counter()
    while i < n:
        comparisons += 1
        if text[i] == pattern[j]:
            i += 1
            j += 1
            if j == m:
                occ += 1
                j = lps[j-1]
        else:
            if j != 0:
                j = lps[j-1]
            else:
                i += 1
    elapsed = (time.perf_counter() - start) * 1000.0
    return occ, comparisons, elapsed

def compare_all():

    # safety - user pressed compare all before analyze
    if "dataset_texts" not in st.session_state or len(st.session_state.dataset_texts)==0:
        st.error("Please analyze CVs first, then click Compare All.")
        return

    # heading at top
    st.markdown("<h2 style='color:#0b7564;'>Compare All Algorithms</h2>", unsafe_allow_html=True)

    all_counts = {}

    for algo_name, algo_func in ALGOS.items():
        algo_counts = {}
        for kw in st.session_state.keywords:
            total = 0
            for txt in st.session_state.dataset_texts:
                occ,_,_ = algo_func(txt.lower(), kw.lower())
                total += occ
            algo_counts[kw] = total
        all_counts[algo_name] = algo_counts

    # performance comparison
    results = []   # rows for table

    for algo_name, algo_func in ALGOS.items():
        total_occ  = 0
        total_time = 0.0
        total_comp = 0

        for txt in st.session_state.dataset_texts:
            for kw in st.session_state.keywords:
                occ, comps, etime = algo_func(txt.lower(), kw.lower())
                total_occ  += occ
                total_time += etime
                total_comp += comps

        results.append([algo_name, total_occ, f"{total_time:.3f}", total_comp])

    # show results table FIRST
    df = pd.DataFrame(results, columns=["Algorithm","Total Occurrences","Exec Time (ms)","Comparisons"])
    st.dataframe(df)

    # show individual keyword occurrences under it
    st.markdown("<h3 style='color:#0b7564;'>Per-Keyword Occurrences</h3>", unsafe_allow_html=True)
    st.dataframe(pd.DataFrame(all_counts))


def word_count_occurrences(text, keyword):
    # case-insensitive strict whole-word match
    pattern = r"\b" + re.escape(keyword.lower()) + r"s?\b"
    return len(re.findall(pattern, text.lower()))

# mapping
ALGOS = {
    "Brute Force": search_naive,
    "Rabin–Karp": search_rabin_karp,
    "KMP": search_kmp
}

JOB_MANDATORY = {
    "Data Analyst": ["python", "sql", "power bi", "excel"],
    "Data Scientist": ["python", "machine learning", "statistics", "numpy"],
    "Frontend Developer": ["html", "css", "javascript", "react"]
}

# PAGE SETUP & CSS
st.set_page_config(page_title="hireLens", layout="wide")
st.markdown("""
<style>

    label, .stTextInput label, .stTextArea label, .stFileUploader label,
    .stMarkdown p {
        color: #111 !important;
    }

    .stAlert, .stAlert p, .stAlert h1, .stAlert h2, .stAlert h3, .stAlert div, .stAlert span {
        color: black !important;
    }
    .stFileUploader div[data-testid="stFileUploaderFileName"] {
        color: black !important;
        opacity: 1.0 !important;
    }

    /* the small “help” text in streamlit inputs */
    .stMarkdown small, .stMarkdown span {
        color: #111 !important;
    }
    [data-testid="stAppViewContainer"] { background-color: white !important; }
    .block-container { padding-left: 6vw !important; padding-right: 6vw !important; }
    .box { background-color: rgba(11,117,100,0.22); padding: 1.2rem; border-radius: 12px; color: black; }
    body { background-color: white !important; color: black; } [data-testid="stAppViewContainer"] { background-color: white !important; } 
    [data-testid="stSidebar"] { background-color: white !important; } /* remove default padding */ div[data-testid="stAppViewContainer"] > .main > div { padding: 0 !important; } /* increased left/right margins for main container */ .block-container { padding-left: 8vw !important; padding-right: 8vw !important; } /* PAGE HEADING */
    h1, h1 span { color:#0b7564 !important; font-size:6vw !important; font-weight:900 !important; line-height:1 !important; margin:0 !important; } 
    .subheading { font-size:2vw !important; color:black !important; margin:0 0 -1rem 0 !important; } 
    /* DESCRIPTION BOX */ .box { background-color: rgba(11,117,100,0.4); padding: 2vw; border-radius: 2vw; font-size: 1.2vw !important; line-height: 1.5; margin-top: 1vw; margin-bottom: 2vw; text-align: justify; color: black; } 
    /* FLEX CONTAINER FOR HEADING + BUTTON */ 
    .header-row { display: flex; flex-wrap: wrap; justify-content: space-between; align-items: center; width: 100%; margin-bottom: 1rem !important; } 
    .header-row button { background-color: #0b7564 !important; color: white !important; border: none; border-radius: 1vw !important; font-size: 2vw !important; padding: 0.8vw 2vw !important; cursor: pointer; } 
    /* IMAGES */ .responsive-img { max-width: 100%; height: auto; } /* ELEMENT1.PNG NEAR DESCRIPTION BOX */ .element1-img { width: 15%; max-width: 200px; margin-left: 2vw; } /* MEDIA QUERIES FOR SMALL SCREENS */ @media (max-width: 1024px) { h1 { font-size: 10vw !important; } .subheading { font-size: 4vw !important; } .box { font-size: 3vw !important; padding: 4vw; border-radius: 4vw; margin-bottom: 3vw; } .header-row button { font-size: 4vw !important; padding: 2vw 4vw !important; } .element1-img { width: 25%; margin-left: 0; display: block; margin-top: 2vw; } } /* sparkle next to description box */ .element1-img { width: 20%; /* adjusts size relative to container */ max-width: 400px; margin-left: 6.5vw; /* spacing from description box */ } /* small screens */ @media (max-width: 1024px){ .element1-img { width: 25%; margin-left: 0; display: block; margin-top: 2vw; /* moves below box on small screens */ } }
</style>
""", unsafe_allow_html=True)

# SESSION STATE
if "page" not in st.session_state:
    st.session_state.page = "home"
if "uploaded_file" not in st.session_state:
    st.session_state.uploaded_file = None
if "uploaded_text" not in st.session_state:
    st.session_state.uploaded_text = None
if "keywords" not in st.session_state:
    st.session_state.keywords = []
if "selected_algo" not in st.session_state:
    st.session_state.selected_algo = "Brute Force"

def go_to(p): st.session_state.page = p

# HOME PAGE
if st.session_state.page == "home":
    # LOAD IMAGES
    logo = Image.open("media/logo.png")
    circle = Image.open("media/element2.png")
    element1 = Image.open("media/element1.png")
    side_img = Image.open("media/image1.png")
    pdf_icon = Image.open("media/pdf.png")
    zip_icon = Image.open("media/zip.png")

    col1, col2 = st.columns([3,1])  # Adjust widths for heading + button layout

    with col1:
        # Logo (small top margin)
        st.markdown(f"""
        <div style="margin-top:-6vw;">
            <img class='responsive-img' src="data:image/png;base64,{img_to_base64('media/logo.png')}" width="150">
        </div>
        """, unsafe_allow_html=True)

        # Heading + circle behind heading
        st.markdown(f"""
        <div style="position:relative; margin-top:1vw;">
            <h1 style="display:inline-block; color:#0b7564; margin:0; z-index:2; position:relative;">hireLens</h1>
            <img src="data:image/png;base64,{img_to_base64('media/element2.png')}"
                style="position:absolute; top:-100%; left:-10%; width:20vw; z-index:0;">
        </div>
        """, unsafe_allow_html=True)

        # Subheading
        st.markdown("<div class='subheading' style='margin-top:0.5vw;'>an intelligent CV analyzer</div>", unsafe_allow_html=True)

        # Description box + sparkle side by side
        element1_base64 = img_to_base64("media/element1.png")
        st.markdown(f"""
        <div style="display:flex; flex-wrap:wrap; align-items:flex-start; gap:1vw; width:100%; margin-top:0.5vw;">
            <div class='box' style="flex:1 1 70%;">
                hireLens is an intelligent CV screening platform that automatically extracts relevant skills and experience from resumes.
                It uses string matching and pattern search algorithms to detect whether a candidate is aligned with a job's skill
                requirements. This saves recruiters hours of manual reading and allows faster, fairer hiring decisions based only on
                relevant competencies.
            </div>
            <div style="flex:0 0 auto; align-self:flex-start;">
                <img class='responsive-img' src="data:image/png;base64,{element1_base64}" style="width:150px; max-width:15vw;">
            </div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        # Streamlit "Get Started" button in the same column as side image
        if st.button("Get Started ▶"):
            go_to("upload")  # Navigate to upload page

        # Side image with reduced top margin
        st.markdown(f"""
        <div style="margin-top:1vw;">
            <img class='responsive-img' src="data:image/png;base64,{img_to_base64('media/image1.png')}" 
                style="width:100%; max-width:750px; margin-left:2vw;">
        </div>
        """, unsafe_allow_html=True)

        # PDF/ZIP icons
        st.markdown(f"""
        <div style="display:flex; justify-content:flex-end; gap:1vw; margin-top:1vw;">
            <img class='responsive-img' src="data:image/png;base64,{img_to_base64('media/pdf.png')}" width="80">
            <img class='responsive-img' src="data:image/png;base64,{img_to_base64('media/zip.png')}" width="80">
        </div>
        """, unsafe_allow_html=True)

# UPLOAD PAGE
elif st.session_state.page == "upload":
    st.markdown("""
    <style>
        /* Make text area labels dark (fixes light grey issue) */
        .stTextArea label, .stTextArea > label, 
        [data-testid="stTextArea"] label,
        [data-testid="stTextArea"] p {
            color: #111 !important;
            font-weight: 500 !important;
        }
        
        /* Make markdown headings visible */
        .stMarkdown h3 {
            color: #0b7564 !important;
            font-weight: 600 !important;
            margin-bottom: 0.5rem !important;
        }
        
        /* Reduce top spacing on upload page */
        .block-container {
            padding-top: 1rem !important;
        }
        
        /* Reduce space above title */
        h1 {
            margin-top: 0 !important;
            padding-top: 0 !important;
            font-size: 3.5rem !important;
            margin-bottom: 1rem !important;
        }
    </style>
    """, unsafe_allow_html=True)
    
    st.title("Upload ZIP dataset of CVs")
    st.write("Upload a .zip containing CVs (pdf/docx). Duplicate filenames removed; PDFs preferred over DOCX.")
    uploaded_file = st.file_uploader("Upload a ZIP file", type=["zip"], accept_multiple_files=False)

    if st.button("Back to Home"):
        go_to("home")
    
    # JOB selection
    job_choice = st.selectbox("Select Job Role", list(JOB_MANDATORY.keys()))

    # show mandatory keywords
    mand = JOB_MANDATORY[job_choice]
    st.markdown(f"**Mandatory Keywords for {job_choice}:**<br>{', '.join([m.title() for m in mand])}", unsafe_allow_html=True)

    # optional keywords input (one per line)
    st.markdown("### OPTIONAL keywords (one per line)")
    kw_text = st.text_area("Optional Keywords (one per line)", height=150, placeholder="e.g.\npandas\nseaborn\nscikit-learn")

    # extract optional keywords
    if kw_text:
        optional_kws = [k.strip().lower() for k in kw_text.splitlines() if k.strip()]
    else:
        optional_kws = []

    # algorithm selection
    algo_choice = st.selectbox("Select algorithm", ["Brute Force", "Rabin–Karp", "KMP"], 
                            index=["Brute Force", "Rabin–Karp", "KMP"].index(st.session_state.selected_algo))
    st.session_state.selected_algo = algo_choice

    # FINAL keyword list = mandatory + optional
    st.session_state.keywords = mand + optional_kws
    st.session_state.patterns = build_patterns(st.session_state.keywords)

    # initialize dataset_texts here (not in loop)
    st.session_state.dataset_texts = []

    if uploaded_file:
        # store uploaded file object in session
        st.session_state.uploaded_file = uploaded_file
        # quick feedback (black text)
        st.markdown(f"<p style='color:black;'>{uploaded_file.name} {uploaded_file.size/1_000_000:.1f}MB</p>", unsafe_allow_html=True)


        st.success("File accepted. Now click Analyze CVs to run the selected algorithm on the dataset.")
 
    if st.button("Analyze CVs"):
        if not st.session_state.uploaded_file:
            st.error("Please upload a ZIP file first.")
        elif not st.session_state.keywords:
            st.error("Please enter at least one keyword (one per line).")
        else:
            go_to("analyze")

# ANALYZE PAGE
elif st.session_state.page == "analyze":
    st.markdown("""
    <style>
        .block-container {
            padding-top: 1rem !important;
        }
        /* Reduce space above title */
        h1 {
            margin-top: 0 !important;
            padding-top: 0 !important;
            font-size: 3.5rem !important;
            margin-bottom: 1rem !important;
        }
    </style>
    """, unsafe_allow_html=True)
    st.title("CV Analysis Result (ZIP dataset)")
    
    # Create two columns for Back button and Compare All button at the top
    col_back, col_compare = st.columns([1, 1])
    
    with col_back:
        if st.button("← Back to Upload"):
            go_to("upload")
    
    with col_compare:
        if st.button("Compare All Algorithms"):
            # Set flag to show comparison
            st.session_state.show_comparison = True
            st.rerun()

    uploaded_file = st.session_state.get("uploaded_file", None)
    
    # Check if we should show comparison view
    if st.session_state.get("show_comparison", False):
        # Reset the flag
        st.session_state.show_comparison = False
        # Call compare_all function
        compare_all()
        st.stop()
    
    # make sure dataset_texts is populated before compare_all
    if "dataset_texts" not in st.session_state or not st.session_state.dataset_texts:
        st.session_state.dataset_texts = []
        # fill it once
        z = zipfile.ZipFile(uploaded_file)
        texts=[]
        for filename in z.namelist():
            if filename.lower().endswith(".pdf") or filename.lower().endswith(".docx"):
                with z.open(filename) as f:
                    if filename.lower().endswith(".pdf"):
                        text = extract_text_from_pdf(f)
                    else:
                        text = extract_text_from_docx(f)
                    texts.append(text)
        st.session_state.dataset_texts = texts

    keywords = st.session_state.get("keywords", [])
    selected_algo = st.session_state.get("selected_algo", "Brute Force")

    if not uploaded_file:
        st.warning("No uploaded ZIP found. Please upload on the Upload page.")
    elif not keywords:
        st.warning("No keywords provided. Go back and add keywords.")
    else:
        st.markdown(f"**Algorithm:** {selected_algo}")
        st.markdown(f"**Keywords:** {', '.join(keywords)}")

        # Open zip and process files (prefer PDFs over docx, remove duplicates)
        z = zipfile.ZipFile(uploaded_file)
        file_names = z.namelist()
        # sort so PDFs come first to prefer them
        file_names.sort(key=lambda x: (not x.lower().endswith(".pdf"), x.lower()))

        processed_base = set()
        results = []  # list of dicts per file
        times_ms = []
        comps_list = []

        search_func = ALGOS[selected_algo]

        for filename in file_names:
            # skip directories
            if filename.endswith('/') or filename.endswith('\\'):
                continue
            base = filename.rsplit('.', 1)[0].lower()
            if base in processed_base:
                continue
            # only consider pdf/docx
            if not (filename.lower().endswith(".pdf") or filename.lower().endswith(".docx")):
                continue

            # open file
            with z.open(filename) as f:
                # read file size from zip info
                try:
                    zinfo = z.getinfo(filename)
                    fsize_bytes = zinfo.file_size
                except Exception:
                    # fallback: read bytes
                    content = f.read()
                    fsize_bytes = len(content)
                    f = io.BytesIO(content)
                # extract text
                if filename.lower().endswith(".pdf"):
                    text = extract_text_from_pdf(f)
                else:
                    text = extract_text_from_docx(f)

            processed_base.add(base)

            # normalize lower for searching
            text_lower = text.lower()

            matched = []
            missing = []
            kw_count_total = 0
            total_time = 0.0
            total_comps = 0

            for kw in keywords:
                kw_clean = kw.strip()
                if not kw_clean: continue

                # strict whole word count (with plural handling)
                occ_count = word_count_occurrences(text, kw_clean)

                # run the selected algorithm for timing/comparisons only
                occ, comps, elapsed = search_func(text_lower, kw_clean.lower())

                total_time += elapsed
                total_comps += comps
                kw_count_total += occ_count

                # store match/missing
                if occ_count > 0:
                    matched.append(f"{kw_clean} ({occ_count})")
                else:
                    missing.append(kw_clean)

            # compute relevance as percentage of keywords matched
            relevance = (len(matched) / len(keywords)) * 100 if keywords else 0.0

            results.append({
                "File": filename,
                "Size (KB)": f"{math.ceil(fsize_bytes/1024)}",
                "Matched Keywords": ", ".join(matched) if matched else "-",
                "Missing Keywords": ", ".join(missing) if missing else "-",
                "Count": kw_count_total,
                "Relevance (%)": round(relevance, 2),
                "Exec Time (ms)": round(total_time, 3),
                "Comparisons": total_comps
            })

            times_ms.append(total_time)
            comps_list.append(total_comps)

        # Prepare DataFrame and sort by relevance desc
        if results:
            df = pd.DataFrame(results)
            df = df.sort_values(by="Relevance (%)", ascending=False).reset_index(drop=True)
            # serial number
            df.index = df.index + 1
            df.index.name = "S.No"

            # Show table styled as HTML so color & text visible
            html_table = df.to_html(index=True, escape=False)

            st.markdown(f"""
            <div style="background-color:#f7fffd;color:black;padding:0.6rem;border-radius:8px;">
                {html_table}
            </div>
            """, unsafe_allow_html=True)

            # Averages
            avg_time = sum(times_ms)/len(times_ms) if times_ms else 0.0
            avg_comps = sum(comps_list)/len(comps_list) if comps_list else 0.0

            st.markdown(f"""
            <div style="margin-top:1rem;padding:0.6rem;background:#eef9f6;border-radius:8px;color:black;">
                <strong>Average execution time (per file):</strong> {avg_time:.3f} ms &nbsp;&nbsp;
                <strong>Average comparisons (per file):</strong> {avg_comps:.1f}
            </div>
            """, unsafe_allow_html=True)

        else:
            st.info("No PDF/DOCX files found in the uploaded ZIP (or all duplicates removed).")